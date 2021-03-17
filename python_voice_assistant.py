import pyttsx3
import datetime
import speech_recognition as sr
import wikipedia
import webbrowser
import os
import random
import pyautogui
from gtts import gTTS

#for resume
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt 

# speech engine initialization
engine = pyttsx3.init('sapi5')
voices = engine.getProperty('voices')
engine.setProperty('voice',voices[1].id)  

#speak function
def speak(audio):
    engine.say(audio)
    engine.runAndWait()

#greeting function
def greetMe():
    hour = int(datetime.datetime.now().hour)
    if hour>=0 and hour<12:
        speak("good morning ma'am. please tell me how can i help you")
    elif hour>=12 and hour<18:
        speak("good Afternoon ma'am. please tell me how can i help you")
    else:
        speak("good night ma'am. please tell me how can i help you")

# take command function to take command from user and convert it to string
def takeCommand():
    r = sr.Recognizer() #initialize the Recognizer
    with sr.Microphone() as source: #use microphone as source of input 
        print("Listening...")
        r.pause_threshold = 1
        audio = r.listen(source)

    try:
        print("Recogniting...")
        query = r.recognize_google(audio , language='en-in') 
        print(f"user said : {query}\n")
    except Exception as e:
        print(e) #print the exception
        print("Please say it again...")
        speak("please say it again.")
        return "None"
    return query



# main function
if  __name__ == "__main__":
    greetMe()
    while True:
       query = takeCommand().lower()
       if 'wikipedia' in query:
           speak("searching wikipedia")
           query = query.replace("wikipedia","")
           result = wikipedia.summary(query , sentences=2)
           speak("According to wikipedia")
           print(result)
           speak(result)

       elif 'play music' in query:
           songs = os.listdir("C:\\Users\\hp\\Desktop\\devi mandir\\कृष्ण भजन\\dinesh")
           print(songs)
           os.startfile(os.path.join("C:\\Users\\hp\\Desktop\\devi mandir\\कृष्ण भजन\\dinesh" , songs[random.randrange(0 , len(songs) - 1)]))
           #used random module for playing random song
        
       elif 'the time' in query:
            curtime = datetime.datetime.now().strftime("%H:%M:%S")
            speak(f"the time is {curtime}")
            print('The current time is '+curtime)

       elif 'open youtube' in query:
           webbrowser.open('youtube.com')

       elif 'open google' in query:
           webbrowser.open('google.com')

       elif 'open stack overflow' in query:
           webbrowser.open('stackoverflow.com')

       elif 'open gmail' in query:
            webbrowser.open_new_tab("gmail.com")
        
       elif 'open code' in query or 'open visual studio' in query:
            os.startfile("C:\\Users\\hp\\AppData\\Local\\Programs\\Microsoft VS Code\\Code.exe")
        
       elif 'open chrome' in query:
            os.startfile("C:\\Program Files (x86)\\Google\\Chrome\\Application\\chrome.exe")

       elif 'open text' in query or 'open sublime' in query:
           os.startfile("C:\\Program Files\\Sublime Text 3\\sublime_text.exe")

       elif 'search' in query:
           query = query.replace("search","")
           webbrowser.open_new_tab(query)

       elif 'in youtube' in query:
           query = query.replace("in youtube" , "")
           webbrowser.open_new_tab('https://www.youtube.com/results?search_query=' + query)
    
       elif 'take a screenshot' in query:
           myss = pyautogui.screenshot()
           myss.save('img.png')
           speak("screenshot is saved , do you want to see results ?")
           ans = takeCommand().lower()
           if 'yes' in ans:
                os.startfile('C:\\Users\\hp\\Desktop\\python\\img.png')
       
       elif 'google map' in query or 'map' in query:
           speak("which places map you want to see")
           place = takeCommand().lower()
           webbrowser.open_new_tab('https://www.google.com/maps/place/'+place)

       elif 'who are you' in query or 'tell me about you' in query or 'what can you do' in query:
           speak("I'm kukky version 1 point o. I can open google, youtube, chrome for you. i can play music if you feel down. i can open your favourite programs. i will greet you everytime you open me")

       elif 'thank you' in query or 'thankyou' in query:
           speak("it's my pleasure")
        
       elif 'who made you' in query or 'who built you' in query:
           speak("I was built by Ritika and Rahul") 
       
       elif "good bye" in query or "ok bye" in query or "stop" in query or 'bye' in query:
            print('your personal assistant kukky is shutting down. have a nice day. Bye')
            speak('your personal assistant kukky is shutting down. have a nice day. Bye')
            break

       elif "resume" in query:
            # -----Initialisation-------
            document = Document() 
             
            # <------- x -------- x --------- x ------- x ------- x ------>

            #style
            run = document.styles['Normal']
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(10) 
            # <------- x -------- x --------- x ------- x ------- x ------>

            # ------Header--------
            speak("Enter your name") 
            para = document.add_paragraph().add_run(input("Enter Your Name : "))
            font = para.font
            font.size = Pt(14)
            font.bold = True
            print("\n\n")

            # <------- x -------- x --------- x ------- x ------- x ------>


            # ------contact----------

            document.add_heading("CONTACT DETAILS")
            speak("Enter your email")
            document.add_paragraph("E-mail : " + input("Enter Your E-mail : "))
            speak("Enter your contact number")
            document.add_paragraph("Contact No. : " + input("Enter Your Contact Number : "))
            speak("Enter Your LinkedIn I'd : ")
            document.add_paragraph("LinkedIn : " + input("Enter Your LinkedIn I'd : "))
            print("\n\n")
            
            # <------- x -------- x --------- x ------- x ------- x ------>


            #<------------- AIM --------------->
            print('\n\n')
            speak("Enter your Objective")
            document.add_heading("Objective : ", level = 1)
            document.add_paragraph(input("Enter your Objective : "))

            # <------- x -------- x --------- x ------- x ------- x ------>


            #<------------- Educational Qualification ------------->

            document.add_heading('Educational Qualification : ', level=1)
            dict = { 1 : "Matric", 2 :"Inter"  , 3 : "Bachelor's" , 4 : "Master's" }
            speak("Please select your Highest Qualification")
            print("\nPlease select your Highest Qualification :")
            for i in range(4):
                strn ="Enter " + str(i+1) + " for " + dict[i+1]
                print(strn)


             # <---- Qualification Table ---->

            n = int(input())
            Qualification_Table = document.add_table( rows = n , cols = 5 )
            header = Qualification_Table.rows[0].cells
            edu_header = {0 :"Education Qualification", 1 :"Institute Name", 2 :"Board/University" , 3 :"Passing Year" ,4 :"Aggregated Marks(Percentage/CGPA)"}
            for i  in range(5) :
                header[i].text = edu_header[i]

             # <---- Qualification Details ----->

            i = 0
            while n > 0 :
                print("\n\n")
                details = Qualification_Table.rows[i].cells
                print(dict[n])
                details[0].text = dict[n]
                speak("Enter the Institute Name")
                details[1].text = input("Enter the Institute Name (Specialization) : ")
                speak("Enter the Board or University")
                details[2].text = input("Enter the Board/University : ")
                speak("Enter the Passing Year")
                details[3].text = input("Enter the Passing Year : ")
                speak("Enter marks aggregated")
                details[4].text = input("Enter marks aggregated : ")
                n -= 1
                i += 1

            # <------- x -------- x --------- x ------- x ------- x ------>


            # -------Skills---------
            print("\n\n")
            document.add_heading("SKILLS")
            document.add_heading("Programming Languages : " , level= 2)
            speak("Enter languages you're expert in")
            document.add_paragraph(input("Enter languages you're expert in : "))

            document.add_heading("IDE",level =2)
            speak("Enter IDEs you have worked with")
            document.add_paragraph(input("Enter IDEs you have worked with : "))

            document.add_heading("Scripting languages : " , level =2)
            speak("Enter Scripting languages you know")
            document.add_paragraph(input("Enter Scripting languages you know : "))

            document.add_heading("Technologies : " , level=2 )
            speak("Enter technologies you know")
            document.add_paragraph(input("Enter technologies you know : "))
            
             # <------- x -------- x --------- x ------- x ------- x ------>


            #--------Experience-------
            print("\n\n")
            speak("Do you want to add any Experience")
            inp = input("Do you want to add any Experience ?")
            inp.lower()
            if 'yes' in inp:
                document.add_heading("EXPERIENCE")
                while True:
                    speak("Enter post in which you worked")
                    document.add_heading(input("Enter post in which you worked : "),level =2)
                    speak("Enter time period and company name")
                    document.add_heading(input("Enter time period and company name : ") , level =3)
                    speak("Enter job description")
                    para1 = document.add_paragraph(input("Enter job description : "), style = 'List Bullet').paragraph_format.left_indent = Inches(0.5)
                    speak("Do you want to add more experience")
                    if input("Do you want to add more experience ?\n ( yes / no ) : ").lower() == "no":
                        break

            # <------- x -------- x --------- x ------- x ------- x ------>


            # --------Projects----------
            print("\n\n")
            speak("Do you want to add any Project")
            inp = input("Do you want to add any project ?")
            inp.lower()
            if 'yes' in inp:
                document.add_heading("Project")
                while True:
                    speak("Enter project in which you worked")
                    document.add_heading(input("Enter project name in which you worked : "),level =2)
                    speak("Enter time period")
                    document.add_heading(input("Enter time period : "), level =3)
                    speak("Enter project description")
                    para1 = document.add_paragraph(input("Enter project  description : ") , style = 'List Bullet').paragraph_format.left_indent = Inches(0.5)
                    speak("Enter project link")
                    para2 = document.add_paragraph("Link :" + input("Enter project link : ")).paragraph_format.left_indent = Inches(0.5)
                    speak("Do you want to add more project")
                    if input("Do you want to add more project ?\n ( yes / no ) : ").lower() == "no":
                        break

            # <------- x -------- x --------- x ------- x ------- x ------>


            # ------Achivements---------
            print("\n\n")
            speak("Do you want to add  any achivement")
            inp = input("Do you want to add any achivement ?")
            inp.lower()
            if 'yes' in inp:
                document.add_heading("ACHIVEMENTS")
                while True:
                    speak("Enter your achivement")
                    document.add_paragraph(input("Enter your achivement :") , style = 'List Bullet')
                    speak("Do you want to add more achivements")
                    if input("Do you want to add more achivements ? \n ( yes / no ) : ").lower() == "no":
                        break
            document.save("Resume.docx")
            
            # <------- x -------- x --------- x ------- x ------- x ------>

            speak("Do you want to overview your resume ?")
            if input("Do you want to overview your resume ?").lower() == "yes":
                os.startfile('C:\\Users\\hp\\Desktop\\python\\Resume.docx')
             